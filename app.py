import streamlit as st
from phi.agent import Agent
from phi.model.google import Gemini
from phi.tools.duckduckgo import DuckDuckGo
from google.generativeai import upload_file, get_file
import google.generativeai as genai

import time
from pathlib import Path
import tempfile
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
load_dotenv()
import os

API_KEY = os.getenv("GOOGLE_API_KEY")
if API_KEY:
    genai.configure(api_key=API_KEY)

st.set_page_config(
    page_title="Meeting Summarizer and Minutes Maker",
    page_icon="🔄",
    layout="wide"
)

st.title("Meeting Summarizer and Minutes Maker 🔄")
st.header("Powered by Gemini 2.0")

@st.cache_resource
def initialize_agent():
    return Agent(
        name="Meeting Summarizer",
        model=Gemini(id="gemini-2.0-flash-exp"),
        tools=[DuckDuckGo()],
        markdown=True,
    )

multimodal_agent = initialize_agent()

uploaded_file = st.file_uploader(
    "Upload a meeting audio or video file", type=['mp4', 'mov', 'avi', 'mp3', 'wav'],
    help="Upload a meeting recording for summarization and minutes generation."
)

def format_text_in_doc(doc, content, title):
    doc.add_heading(title, level=1)
    paragraphs = content.split("\n")
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = p.runs[0]
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0, 0, 0)

def generate_document(response_content, file_name, title):
    doc_file = Path(tempfile.gettempdir()) / file_name
    doc = Document()
    format_text_in_doc(doc, response_content, title)
    doc.save(doc_file)
    return doc_file

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as temp_file:
        temp_file.write(uploaded_file.read())
        file_path = temp_file.name

    if uploaded_file.type.startswith("video"):
        st.video(file_path, format="video/mp4", start_time=0)
    elif uploaded_file.type.startswith("audio"):
        st.audio(file_path, format="audio/mp3")

    st.subheader("Choose an Option")

    col1, col2, col3, col4, col5 = st.columns(5)

    if col1.button("✍️ Generate Minutes"):
        try:
            with st.spinner("Generating meeting minutes..."):
                processed_file = upload_file(file_path)
                while processed_file.state.name == "PROCESSING":
                    time.sleep(1)
                    processed_file = get_file(processed_file.name)

                minutes_prompt = (
                    """
                    Generate detailed and structured minutes from the uploaded meeting recording. Include
                    key points, decisions made, and action items.
                    """
                )

                response = multimodal_agent.run(minutes_prompt, videos=[processed_file])

            st.subheader("Meeting Minutes")
            st.markdown(response.content)

            doc_file = generate_document(response.content, "meeting_minutes.docx", "Meeting Minutes")

            st.download_button(
                label="🔗 Download Minutes",
                data=open(doc_file, "rb").read(),
                file_name="meeting_minutes.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as error:
            st.error(f"An error occurred: {error}")

    if col2.button("🔍 Generate Summary"):
        try:
            with st.spinner("Generating meeting summary..."):
                processed_file = upload_file(file_path)
                while processed_file.state.name == "PROCESSING":
                    time.sleep(1)
                    processed_file = get_file(processed_file.name)

                summary_prompt = (
                    """
                    Summarize the uploaded meeting recording. Include the main topics discussed,
                    major takeaways, and any conclusions reached.
                    """
                )

                response = multimodal_agent.run(summary_prompt, videos=[processed_file])

            st.subheader("Meeting Summary")
            st.markdown(response.content)

            doc_file = generate_document(response.content, "meeting_summary.docx", "Meeting Summary")

            st.download_button(
                label="🔗 Download Summary",
                data=open(doc_file, "rb").read(),
                file_name="meeting_summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as error:
            st.error(f"An error occurred: {error}")

    if col3.button("🗒️ Action Items"):
        try:
            with st.spinner("Extracting action items..."):
                processed_file = upload_file(file_path)
                while processed_file.state.name == "PROCESSING":
                    time.sleep(1)
                    processed_file = get_file(processed_file.name)

                action_items_prompt = (
                    """
                    Extract all actionable items discussed in the meeting. Provide a list of tasks, responsible parties, and deadlines if mentioned.
                    """
                )

                response = multimodal_agent.run(action_items_prompt, videos=[processed_file])

            st.subheader("Action Items")
            st.markdown(response.content)

            doc_file = generate_document(response.content, "action_items.docx", "Action Items")

            st.download_button(
                label="🔗 Download Action Items",
                data=open(doc_file, "rb").read(),
                file_name="action_items.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as error:
            st.error(f"An error occurred: {error}")

    if col4.button("📊 Generate Insights"):
        try:
            with st.spinner("Generating insights..."):
                processed_file = upload_file(file_path)
                while processed_file.state.name == "PROCESSING":
                    time.sleep(1)
                    processed_file = get_file(processed_file.name)

                insights_prompt = (
                    """
                    Provide high-level insights and analytics based on the meeting. Highlight key trends, concerns, and positive outcomes.
                    """
                )

                response = multimodal_agent.run(insights_prompt, videos=[processed_file])

            st.subheader("Meeting Insights")
            st.markdown(response.content)

            doc_file = generate_document(response.content, "meeting_insights.docx", "Meeting Insights")

            st.download_button(
                label="🔗 Download Insights",
                data=open(doc_file, "rb").read(),
                file_name="meeting_insights.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as error:
            st.error(f"An error occurred: {error}")

# Sidebar for Chat with Meeting
if uploaded_file:
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []

    with st.expander("💬 Chat with Meeting", expanded=True):
        st.title("Chat with Meeting")

        query = st.text_input("Ask a question about the meeting:", key="sidebar_query")

        if st.button("Send Query", key="sidebar_send"):
            if query.strip() != "":
                try:
                    with st.spinner("Generating response..."):
                        processed_file = upload_file(file_path)
                        while processed_file.state.name == "PROCESSING":
                            time.sleep(1)
                            processed_file = get_file(processed_file.name)

                        chat_prompt = (
                            f"""
                            Analyse the video meeting and answer the users questions concisely.

                            Question: {query}
                            """
                        )
                        chat_response = multimodal_agent.run(chat_prompt, videos=[processed_file])
                        st.session_state.chat_history.append((query, chat_response.content))

                except Exception as error:
                    st.error(f"An error occurred during the chat: {error}")

        if st.session_state.chat_history:
            st.subheader("Chat History")
            for i, (q, a) in enumerate(st.session_state.chat_history):
                st.markdown(f"**Q{i+1}:** {q}")
                st.markdown(f"**A{i+1}:** {a}")

else:
    st.info("Upload a meeting recording to begin analysis.")

st.markdown(
    """
    <style>
    .stTextArea textarea {
        height: 100px;
    }
    </style>
    """,
    unsafe_allow_html=True
)
